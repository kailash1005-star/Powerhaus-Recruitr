"""
Document Parsing Service — document bytes → clean text / Markdown.

Lightweight, dependency-slim parsers:
  * PDF  → pypdf (text layer extraction)
  * DOCX → python-docx (paragraphs + tables)
  * HTML → markdownify (falls back to BeautifulSoup text)
  * TXT/MD/CSV and anything decodable → decoded directly

This REPLACED Docling. Docling's only job here was this same parsing, but it
pulled the entire PyTorch stack (torch + torchvision + CUDA wheels + model
weights) — a multi-GB image and ~18-20 min builds — for a feature the rest of
the app never touches. These parsers are pure-Python and add a few hundred KB.

Trade-off (accepted): no ML layout/table analysis and no OCR of scanned/image
PDFs. Digital, text-based CVs/JDs — the overwhelming norm — parse fine. A scanned
(image-only) PDF yields no text layer; the caller already treats empty output as
a failed ingestion ("no text extracted"), so it's flagged rather than silently
wrong.

Design notes:
  * Synchronous & CPU-bound; callers use `parse_bytes`, which offloads to a
    worker thread so the FastAPI event loop is never blocked.
"""
from __future__ import annotations

import asyncio
import io
import logging
import os

logger = logging.getLogger(__name__)

_TEXT_EXTS = {".txt", ".md", ".markdown", ".text", ".csv"}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
_HTML_EXTS = {".html", ".htm"}


def _ext(filename: str | None) -> str:
    return os.path.splitext(filename or "")[1].lower()


def _parse_pdf(data: bytes) -> str:
    """Extract the text layer of a digital PDF, page by page."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001 — one bad page shouldn't kill the doc
            logger.warning("[parse] PDF page extraction failed: %s", e)
            txt = ""
        if txt.strip():
            parts.append(txt.strip())
    return "\n\n".join(parts).strip()


def _parse_docx(data: bytes) -> str:
    """Extract paragraphs and tables from a .docx (Office Open XML)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    # Render tables as pipe-joined rows so skills/experience grids stay readable.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def _parse_html(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    try:
        from markdownify import markdownify as md

        return md(text).strip()
    except Exception:  # noqa: BLE001 — fall back to plain text extraction
        from bs4 import BeautifulSoup

        return BeautifulSoup(text, "html.parser").get_text("\n").strip()


def _parse_sync(data: bytes, filename: str | None) -> str:
    ext = _ext(filename)

    if ext in _TEXT_EXTS:
        return data.decode("utf-8", errors="replace").strip()
    if ext in _PDF_EXTS:
        return _parse_pdf(data)
    if ext in _DOCX_EXTS:
        return _parse_docx(data)
    if ext in _HTML_EXTS:
        return _parse_html(data)

    # Unknown/missing extension — sniff magic bytes, then fall back to decoding.
    if data[:5] == b"%PDF-":
        return _parse_pdf(data)
    if data[:2] == b"PK":  # zip container → almost certainly a .docx
        try:
            return _parse_docx(data)
        except Exception:  # noqa: BLE001
            pass
    try:
        text = data.decode("utf-8")
        if text.strip():
            return text.strip()
    except UnicodeDecodeError:
        pass

    raise ValueError(
        f"Unsupported document type '{ext or 'unknown'}'. Supported: PDF, DOCX, "
        "HTML, and plain text. Scanned-image PDFs and legacy .doc are not "
        "supported (no OCR/converter in this build)."
    )


async def parse_bytes(data: bytes, filename: str | None = None) -> str:
    """Async entry point — offloads the blocking parse to a worker thread.

    Returns clean text/Markdown. Raises on unsupported type or empty input;
    callers treat empty output as a failed ingestion.
    """
    if not data:
        raise ValueError("empty document")
    return await asyncio.to_thread(_parse_sync, data, filename)


def warm_up() -> None:
    """No-op, kept for call-site compatibility (no heavy models to preload)."""
    return None
